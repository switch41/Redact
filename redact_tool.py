import spacy
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from PIL import Image, ImageDraw, UnidentifiedImageError
import pytesseract
from pdf2image import convert_from_path
from transformers import AutoModelForTokenClassification, AutoTokenizer
import torch
import docx
import openpyxl
import threading
import os
import json
import logging
from ttkthemes import ThemedTk
from datetime import datetime

# Configure logging
logging.basicConfig(
    filename='redaction.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Set Tesseract path to current directory
current_dir = os.path.dirname(os.path.abspath(__file__))
pytesseract.pytesseract.tesseract_cmd = os.path.join(current_dir, 'tesseract.exe')

# Load spaCy model for NER
try:
    nlp = spacy.load("en_core_web_sm")
except Exception as e:
    logging.error(f"Error loading spaCy model: {str(e)}")
    messagebox.showerror("Error", "Failed to load language model. Please ensure spaCy is properly installed.")
    exit(1)

# Global variables
transformer_model = None
tokenizer = None
redaction_history = []
current_history_index = -1
settings = {}
custom_patterns = {}

# Enhanced security patterns
DEFAULT_PATTERNS = {
    'emails': r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
    'phone_numbers': r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
    'ssn': r'\b\d{3}[-]?\d{2}[-]?\d{4}\b',
    'credit_cards': r'\b\d{4}[-]?\d{4}[-]?\d{4}[-]?\d{4}\b',
    'bank_accounts': r'\b\d{8,12}\b',
    'addresses': r'\b\d+\s+[A-Za-z\s,]+\b[A-Z]{2}\s+\d{5}\b',
    'dates': r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b',
    'ip_addresses': r'\b\d{1,3}\.\d{1,3}\.\d{1,3}\.\d{1,3}\b',
    'urls': r'https?://(?:[-\w.]|(?:%[\da-fA-F]{2}))+[^\s]*'
}

class RedactionHistory:
    def __init__(self):
        self.history = []
        self.current = -1
    
    def add_state(self, state):
        self.history = self.history[:self.current + 1]
        self.history.append(state)
        self.current += 1
    
    def undo(self):
        if self.current > 0:
            self.current -= 1
            return self.history[self.current]
        return None
    
    def redo(self):
        if self.current < len(self.history) - 1:
            self.current += 1
            return self.history[self.current]
        return None

class CustomRule:
    def __init__(self):
        self.patterns = []
        self.keywords = []
    
    def add_pattern(self, pattern):
        try:
            re.compile(pattern)
            self.patterns.append(pattern)
            return True
        except re.error:
            return False
    
    def add_keyword(self, keyword):
        self.keywords.append(keyword.lower())

def save_settings():
    settings = {
        'redaction_level': redaction_var.get(),
        'custom_patterns': custom_patterns,
        'last_directory': os.path.dirname(entry_file_path.get()) if entry_file_path.get() else "",
        'theme': root.get_theme()
    }
    try:
        with open('redact_settings.json', 'w') as f:
            json.dump(settings, f)
        logging.info("Settings saved successfully")
    except Exception as e:
        logging.error(f"Error saving settings: {str(e)}")

def load_settings():
    try:
        with open('redact_settings.json', 'r') as f:
            return json.load(f)
    except FileNotFoundError:
        return {}
    except Exception as e:
        logging.error(f"Error loading settings: {str(e)}")
        return {}

def update_progress(value):
    progress_bar['value'] = value
    root.update_idletasks()

def preview_redaction():
    input_file = entry_file_path.get()
    if not input_file:
        messagebox.showerror("Error", "Please select a file first.")
        return
    
    redaction_level = redaction_var.get()
    
    preview_window = tk.Toplevel(root)
    preview_window.title("Preview Redacted Text")
    preview_window.geometry("600x400")
    
    preview_frame = ttk.Frame(preview_window, padding="10")
    preview_frame.pack(fill=tk.BOTH, expand=True)
    
    preview_text = tk.Text(preview_frame, wrap=tk.WORD)
    preview_text.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    scrollbar = ttk.Scrollbar(preview_frame, orient=tk.VERTICAL, command=preview_text.yview)
    scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
    preview_text['yscrollcommand'] = scrollbar.set
    
    try:
        if input_file.endswith('.txt'):
            with open(input_file, 'r', encoding='utf-8') as f:
                text = f.read()
        elif input_file.endswith('.docx'):
            doc = docx.Document(input_file)
            text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        else:
            messagebox.showinfo("Info", "Preview is only available for text and Word documents.")
            return
        
        redacted = redact_text_with_ner(text, redaction_level)
        preview_text.insert('1.0', redacted)
        preview_text.config(state='disabled')
        
    except Exception as e:
        logging.error(f"Error in preview: {str(e)}")
        messagebox.showerror("Error", f"Failed to generate preview: {str(e)}")

def batch_process():
    folder_path = filedialog.askdirectory(title="Select Folder for Batch Processing")
    if not folder_path:
        return
    
    output_folder = os.path.join(folder_path, "redacted_files")
    os.makedirs(output_folder, exist_ok=True)
    
    files = [f for f in os.listdir(folder_path) 
             if f.endswith(('.txt', '.docx', '.pdf', '.xlsx', '.jpg', '.jpeg', '.png'))]
    
    if not files:
        messagebox.showinfo("Info", "No supported files found in the selected folder.")
        return
    
    progress_window = tk.Toplevel(root)
    progress_window.title("Batch Processing")
    progress_window.geometry("300x150")
    
    progress_label = ttk.Label(progress_window, text="Processing files...")
    progress_label.pack(pady=10)
    
    batch_progress = ttk.Progressbar(progress_window, length=200, mode='determinate')
    batch_progress.pack(pady=10)
    
    for i, file in enumerate(files):
        try:
            input_path = os.path.join(folder_path, file)
            output_path = os.path.join(output_folder, f"redacted_{file}")
            
            if file.endswith('.txt'):
                redact_file_with_ner(input_path, output_path, redaction_var.get())
            elif file.endswith('.docx'):
                redact_word_file(input_path, output_path, redaction_var.get())
            elif file.endswith('.pdf'):
                redact_pdf(input_path, output_path, redaction_var.get())
            elif file.endswith('.xlsx'):
                redact_excel_file(input_path, output_path, redaction_var.get())
            elif file.endswith(('.jpg', '.jpeg', '.png')):
                redact_image(input_path, output_path, redaction_var.get())
                
            batch_progress['value'] = ((i + 1) / len(files)) * 100
            progress_window.update()
            
        except Exception as e:
            logging.error(f"Error processing {file}: {str(e)}")
    
    progress_window.destroy()
    messagebox.showinfo("Complete", f"Batch processing complete. Files saved in:\n{output_folder}")

def toggle_theme():
    current_theme = root.get_theme()
    new_theme = 'equilux' if current_theme == 'arc' else 'arc'
    root.set_theme(new_theme)
    save_settings()

def create_custom_pattern():
    pattern_window = tk.Toplevel(root)
    pattern_window.title("Add Custom Pattern")
    pattern_window.geometry("400x200")
    
    ttk.Label(pattern_window, text="Pattern Name:").pack(pady=5)
    name_entry = ttk.Entry(pattern_window)
    name_entry.pack(pady=5)
    
    ttk.Label(pattern_window, text="Regular Expression:").pack(pady=5)
    pattern_entry = ttk.Entry(pattern_window)
    pattern_entry.pack(pady=5)
    
    def save_pattern():
        name = name_entry.get()
        pattern = pattern_entry.get()
        try:
            re.compile(pattern)
            custom_patterns[name] = pattern
            pattern_window.destroy()
            messagebox.showinfo("Success", "Custom pattern added successfully!")
        except re.error:
            messagebox.showerror("Error", "Invalid regular expression pattern!")
    
    ttk.Button(pattern_window, text="Save Pattern", command=save_pattern).pack(pady=20)

def redact_text_with_ner(text, redaction_level=1):
    doc = nlp(text)
    redacted_text = text

    # Apply NER-based redaction
    for ent in doc.ents:
        if redaction_level == 1 and ent.label_ in ['PERSON']:
            redacted_text = redacted_text.replace(ent.text, '[REDACTED]')
        elif redaction_level == 2 and ent.label_ in ['PERSON', 'ORG', 'GPE']:
            redacted_text = redacted_text.replace(ent.text, '[REDACTED]')
        elif redaction_level == 3:
            redacted_text = redacted_text.replace(ent.text, '[REDACTED]')

    # Apply pattern-based redactions
    patterns_to_use = DEFAULT_PATTERNS.copy()
    if custom_patterns:
        patterns_to_use.update(custom_patterns)

    if redaction_level >= 2:
        for pattern_name, pattern in patterns_to_use.items():
            redacted_text = re.sub(pattern, '[REDACTED]', redacted_text)

    return redacted_text

def select_file():
    file_path = filedialog.askopenfilename(
        title="Select a File",
        filetypes=(
            ("Text Files", "*.txt"),
            ("Word Files", "*.docx"),
            ("Excel Files", "*.xlsx"),
            ("PDF Files", "*.pdf"),
            ("Image Files", "*.jpg;*.jpeg;*.png;*.bmp;*.gif;*.tiff"),
            ("All Files", "*.*")
        )
    )
    if file_path:
        entry_file_path.delete(0, tk.END)
        entry_file_path.insert(0, file_path)
        status_bar.config(text=f"Selected file: {os.path.basename(file_path)}")

def redact_file_with_ner(input_file, output_file, redaction_level=1):
    try:
        with open(input_file, 'r', encoding='utf-8') as infile:
            text = infile.read()
            redacted_text = redact_text_with_ner(text, redaction_level)
            
        with open(output_file, 'w', encoding='utf-8') as outfile:
            outfile.write(redacted_text)
            
        status_bar.config(text=f"Successfully redacted: {os.path.basename(output_file)}")
        messagebox.showinfo("Success", f"Redacted content saved to {output_file}")
            
    except Exception as e:
        logging.error(f"Error in redact_file_with_ner: {str(e)}")
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def redact_word_file(input_file, output_file, redaction_level=1):
    try:
        doc = docx.Document(input_file)
        for paragraph in doc.paragraphs:
            paragraph.text = redact_text_with_ner(paragraph.text, redaction_level)
        
        doc.save(output_file)
        status_bar.config(text=f"Successfully redacted: {os.path.basename(output_file)}")
        messagebox.showinfo("Success", f"Redacted content saved to {output_file}")

    except Exception as e:
        logging.error(f"Error in redact_word_file: {str(e)}")
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def redact_excel_file(input_file, output_file, redaction_level=1):
    try:
        wb = openpyxl.load_workbook(input_file)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        cell.value = redact_text_with_ner(cell.value, redaction_level)
        
        wb.save(output_file)
        status_bar.config(text=f"Successfully redacted: {os.path.basename(output_file)}")
        messagebox.showinfo("Success", f"Redacted content saved to {output_file}")

    except Exception as e:
        logging.error(f"Error in redact_excel_file: {str(e)}")
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def redact_image(input_image, output_image, redaction_level=1):
    try:
        image = Image.open(input_image)
        gray_image = image.convert('L').resize((image.width // 2, image.height // 2))
        data = pytesseract.image_to_data(gray_image, output_type=pytesseract.Output.DICT)
        draw = ImageDraw.Draw(image)

        for i in range(len(data['text'])):
            if int(data['conf'][i]) > 60:
                text = data['text'][i]
                x, y, w, h = data['left'][i], data['top'][i], data['width'][i], data['height'][i]
                if redaction_level == 1 and nlp(text).ents and nlp(text).ents[0].label_ == 'PERSON':
                    draw.rectangle([x, y, x + w, y + h], fill="black")
                elif redaction_level == 2 and nlp(text).ents and nlp(text).ents[0].label_ in ['PERSON', 'ORG', 'GPE']:
                    draw.rectangle([x, y, x + w, y + h], fill="black")
                elif redaction_level == 3:
                    draw.rectangle([x, y, x + w, y + h], fill="black")

        image.save(output_image)
        status_bar.config(text=f"Successfully redacted: {os.path.basename(output_image)}")
        messagebox.showinfo("Success", f"Redacted image saved to {output_image}")

    except UnidentifiedImageError:
        logging.error("Invalid image format")
        messagebox.showerror("Error", "The selected file is not a valid image format.")
    except Exception as e:
        logging.error(f"Error in redact_image: {str(e)}")
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def redact_pdf(input_pdf, output_pdf, redaction_level=1):
    try:
        pages = convert_from_path(input_pdf)
        total_pages = len(pages)

        with open(output_pdf, 'w', encoding='utf-8') as outfile:
            for page_num, page in enumerate(pages):
                text = pytesseract.image_to_string(page)
                redacted_text = redact_text_with_ner(text, redaction_level)
                outfile.write(f"Page {page_num + 1}\n{'-' * 20}\n")
                outfile.write(redacted_text + "\n\n")
                
                # Update progress
                progress = ((page_num + 1) / total_pages) * 100
                update_progress(progress)

        status_bar.config(text=f"Successfully redacted: {os.path.basename(output_pdf)}")
        messagebox.showinfo("Success", f"Redacted PDF content saved to {output_pdf}")

    except Exception as e:
        logging.error(f"Error in redact_pdf: {str(e)}")
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def run_redaction_in_thread():
    threading.Thread(target=perform_redaction, daemon=True).start()

def perform_redaction():
    input_file = entry_file_path.get()
    redaction_level = redaction_var.get()

    if not input_file:
        messagebox.showerror("Error", "Please select a file.")
        return

    output_file = filedialog.asksaveasfilename(
        title="Save Redacted File As",
        defaultextension=".txt",
        filetypes=(
            ("Text Files", "*.txt"),
            ("Word Files", "*.docx"),
            ("Excel Files", "*.xlsx"),
            ("PDF Files", "*.pdf"),
            ("Image Files", "*.jpg;*.jpeg;*.png;*.bmp;*.gif;*.tiff"),
            ("All Files", "*.*")
        )
    )

    if output_file:
        status_bar.config(text="Processing...")
        progress_bar['value'] = 0
        
        try:
            if input_file.endswith(".txt"):
                redact_file_with_ner(input_file, output_file, redaction_level)
            elif input_file.lower().endswith((".jpg", ".jpeg", ".png", ".bmp", ".gif", ".tiff")):
                redact_image(input_file, output_file, redaction_level)
            elif input_file.endswith(".pdf"):
                redact_pdf(input_file, output_file, redaction_level)
            elif input_file.endswith(".docx"):
                redact_word_file(input_file, output_file, redaction_level)
            elif input_file.endswith(".xlsx"):
                redact_excel_file(input_file, output_file, redaction_level)
                
            progress_bar['value'] = 100
            redaction_history.add_state({
                'input_file': input_file,
                'output_file': output_file,
                'redaction_level': redaction_level
            })
            
        except Exception as e:
            logging.error(f"Error in perform_redaction: {str(e)}")
            messagebox.showerror("Error", f"An error occurred: {str(e)}")
            status_bar.config(text="Error occurred during processing")
            progress_bar['value'] = 0

# Main GUI setup
root = ThemedTk(theme="arc")
root.title("RE-DACT: Advanced Redaction Tool")
root.geometry("800x600")

# Create main frame with padding
main_frame = ttk.Frame(root, padding="20")
main_frame.pack(fill=tk.BOTH, expand=True)

# Menu Bar
menubar = tk.Menu(root)
root.config(menu=menubar)

file_menu = tk.Menu(menubar, tearoff=0)
menubar.add_cascade(label="File", menu=file_menu)
file_menu.add_command(label="Open", command=select_file)
file_menu.add_command(label="Batch Process", command=batch_process)
file_menu.add_separator()
file_menu.add_command(label="Exit", command=root.quit)

edit_menu = tk.Menu(menubar, tearoff=0)
menubar.add_cascade(label="Edit", menu=edit_menu)
edit_menu.add_command(label="Add Custom Pattern", command=create_custom_pattern)
edit_menu.add_command(label="Preview Redaction", command=preview_redaction)

view_menu = tk.Menu(menubar, tearoff=0)
menubar.add_cascade(label="View", menu=view_menu)
view_menu.add_command(label="Toggle Theme", command=toggle_theme)

# File selection frame
file_frame = ttk.LabelFrame(main_frame, text="File Selection", padding="10")
file_frame.pack(fill=tk.X, pady=10)

entry_file_path = ttk.Entry(file_frame, width=50)
entry_file_path.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

btn_browse = ttk.Button(file_frame, text="Browse", command=select_file)
btn_browse.pack(side=tk.RIGHT, padx=5)

# Redaction level frame
level_frame = ttk.LabelFrame(main_frame, text="Redaction Level", padding="10")
level_frame.pack(fill=tk.X, pady=10)

redaction_var = tk.IntVar(value=1)

radio_level_1 = ttk.Radiobutton(level_frame, text="Level 1: Basic (Names)", 
                               variable=redaction_var, value=1)
radio_level_1.pack(anchor="w", padx=20)

radio_level_2 = ttk.Radiobutton(level_frame, text="Level 2: Intermediate (Names, Emails, Locations)", 
                               variable=redaction_var, value=2)
radio_level_2.pack(anchor="w", padx=20)

radio_level_3 = ttk.Radiobutton(level_frame, text="Level 3: Advanced (All Identifiers)", 
                               variable=redaction_var, value=3)
radio_level_3.pack(anchor="w", padx=20)

# Progress frame
progress_frame = ttk.LabelFrame(main_frame, text="Progress", padding="10")
progress_frame.pack(fill=tk.X, pady=10)

progress_bar = ttk.Progressbar(progress_frame, mode='determinate', length=100)
progress_bar.pack(fill=tk.X, padx=5)

# Action buttons frame
button_frame = ttk.Frame(main_frame)
button_frame.pack(fill=tk.X, pady=20)

btn_preview = ttk.Button(button_frame, text="Preview", command=preview_redaction)
btn_preview.pack(side=tk.LEFT, padx=5)

btn_redact = ttk.Button(button_frame, text="Redact", command=run_redaction_in_thread)
btn_redact.pack(side=tk.RIGHT, padx=5)

btn_batch = ttk.Button(button_frame, text="Batch Process", command=batch_process)
btn_batch.pack(side=tk.RIGHT, padx=5)

# Status bar
status_bar = ttk.Label(root, text="Ready", relief=tk.SUNKEN, anchor=tk.W)
status_bar.pack(side=tk.BOTTOM, fill=tk.X)

# Load saved settings
saved_settings = load_settings()
if saved_settings:
    if 'redaction_level' in saved_settings:
        redaction_var.set(saved_settings['redaction_level'])
    if 'theme' in saved_settings:
        root.set_theme(saved_settings['theme'])
    if 'custom_patterns' in saved_settings:
        custom_patterns = saved_settings['custom_patterns']

# Initialize redaction history
redaction_history = RedactionHistory()

root.mainloop()
